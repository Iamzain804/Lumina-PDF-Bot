"""Multi-format document processing for PDF, TXT, MD, and DOCX files."""

import os
from pathlib import Path
from typing import List, Optional

import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Try to import docx for Word documents
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Try to import advanced embeddings, fallback to lightweight
try:
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS
    import sentence_transformers
    ADVANCED_AVAILABLE = True
except ImportError:
    ADVANCED_AVAILABLE = False

# Always import lightweight fallback
from src.engine.lightweight_embeddings import LightweightEmbeddings
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

class DocumentVectorStore:
    """Multi-format document processing and vector database management."""
    
    def __init__(self, config):
        """Initialize embedding model and configuration.
        
        Args:
            config: Configuration object with model settings
        """
        try:
            self.config = config
            
            # Always use lightweight embeddings to avoid dependency issues
            print("Using lightweight TF-IDF embeddings (PyTorch-free)")
            self.embeddings = LightweightEmbeddings()
            self.use_faiss = False
            self.vectors = None
            self.texts = None
            
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=config.CHUNK_SIZE,
                chunk_overlap=config.CHUNK_OVERLAP
            )
        except Exception as e:
            raise RuntimeError(f"Failed to initialize DocumentVectorStore: {e}")
    
    def load_document(self, file_path: str) -> str:
        """Extract text from various document formats.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Full text content from document
        """
        try:
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                return self._load_pdf(file_path)
            elif file_extension in ['.txt', '.md']:
                return self._load_text(file_path)
            elif file_extension == '.docx':
                return self._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise RuntimeError(f"Failed to load document {file_path}: {e}")
    
    def _load_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        if not text.strip():
            raise ValueError("No text found in PDF")
        
        return text
    
    def _load_text(self, file_path: Path) -> str:
        """Extract text from TXT/MD files."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                
                if not text.strip():
                    raise ValueError("No text found in file")
                
                return text
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any encoding")
    
    def _load_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise ValueError("No text found in DOCX file")
        
        return text
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks with metadata.
        
        Args:
            text: Full text content
            
        Returns:
            List of text chunks
        """
        try:
            chunks = self.text_splitter.split_text(text)
            
            if not chunks:
                raise ValueError("No chunks created from text")
            
            return chunks
        
        except Exception as e:
            raise RuntimeError(f"Failed to split text: {e}")
    
    def create_embeddings(self, chunks: List[str]):
        """Create vectorstore from text chunks.
        
        Args:
            chunks: List of text chunks
            
        Returns:
            Lightweight vectorstore dictionary
        """
        try:
            # Always use lightweight embeddings
            vectors = self.embeddings.embed_documents(chunks)
            self.vectors = vectors
            self.texts = chunks
            return {"vectors": vectors, "texts": chunks}
        
        except Exception as e:
            raise RuntimeError(f"Failed to create embeddings: {e}")
    
    def save_vectorstore(self, vectorstore, doc_name: str) -> None:
        """Save vectorstore to disk.
        
        Args:
            vectorstore: Vectorstore dictionary
            doc_name: Name of document file (without extension)
        """
        try:
            save_path = Path(self.config.VECTOR_STORE_DIR) / doc_name
            save_path.mkdir(parents=True, exist_ok=True)
            
            # Save lightweight version
            import pickle
            with open(save_path / "vectors.pkl", "wb") as f:
                pickle.dump(vectorstore, f)
        
        except Exception as e:
            raise RuntimeError(f"Failed to save vectorstore: {e}")
    
    def load_vectorstore(self, doc_name: str):
        """Load existing vectorstore from disk.
        
        Args:
            doc_name: Name of document file (without extension)
            
        Returns:
            Vectorstore dictionary or None if not found
        """
        try:
            load_path = Path(self.config.VECTOR_STORE_DIR) / doc_name
            
            if not load_path.exists():
                return None
            
            # Load lightweight version
            import pickle
            vector_file = load_path / "vectors.pkl"
            if vector_file.exists():
                with open(vector_file, "rb") as f:
                    return pickle.load(f)
            return None
        
        except Exception as e:
            return None
    
    def search(self, query: str, vectorstore, k: int) -> List[tuple]:
        """Perform similarity search in vectorstore.
        
        Args:
            query: Search query
            vectorstore: Vectorstore dictionary
            k: Number of results to return
            
        Returns:
            List of (document, score) tuples
        """
        try:
            # Lightweight search
            query_vector = self.embeddings.embed_query(query)
            vectors = vectorstore["vectors"]
            texts = vectorstore["texts"]
            
            similarities = cosine_similarity([query_vector], vectors)[0]
            top_indices = np.argsort(similarities)[::-1][:k]
            
            results = []
            for idx in top_indices:
                # Create document-like object
                doc = type('Document', (), {'page_content': texts[idx]})()
                score = float(similarities[idx])
                results.append((doc, score))
            
            return results
        
        except Exception as e:
            raise RuntimeError(f"Failed to search vectorstore: {e}")